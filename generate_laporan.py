from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins: Left 1.25", Right/Top/Bottom 1.0", Header/Footer 0.5" ──
section = doc.sections[0]
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

# ── Helper: set paragraph font ──
def set_font(paragraph, size=12, bold=False, italic=False, color=None):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

def set_para_format(paragraph, spacing=1.5, align='justify', tab=False, space_before=0, space_after=6):
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(spacing * 12)
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if align == 'justify':
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if tab:
        pf.first_line_indent = Cm(1.25)

def add_body(text, tab=False, bold_parts=None):
    """Add justified body paragraph, optional tab indent, optional bold spans."""
    p = doc.add_paragraph()
    set_para_format(p, spacing=1.5, align='justify', tab=tab, space_after=0)
    if bold_parts:
        # bold_parts = list of (substring, is_bold)
        for chunk, is_bold in bold_parts:
            run = p.add_run(chunk)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = is_bold
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_heading_bab(bab_text, title_text):
    """BAB I / PENGENALAN style headings."""
    doc.add_paragraph()  # spacing before
    p_bab = doc.add_paragraph()
    set_para_format(p_bab, align='center', space_before=24, space_after=0)
    r = p_bab.add_run(bab_text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

    p_title = doc.add_paragraph()
    set_para_format(p_title, align='center', space_before=0, space_after=18)
    r2 = p_title.add_run(title_text)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.bold = True

def add_section_heading(number, title):
    """1.1 PENGENALAN PROJEK style."""
    p = doc.add_paragraph()
    set_para_format(p, align='left', space_before=12, space_after=6)
    r = p.add_run(f"{number}\t{title}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

def add_subheading(text):
    """Bold sub-heading like Fasa 1: ..."""
    p = doc.add_paragraph()
    set_para_format(p, align='left', space_before=8, space_after=0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

def add_blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)

def add_list_item(prefix, text):
    p = doc.add_paragraph()
    set_para_format(p, align='justify', space_before=0, space_after=0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f"{prefix}\t{text}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

# ════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════
p = doc.add_paragraph()
set_para_format(p, align='center', space_before=72, space_after=0)
r = p.add_run("LAPORAN PROJEK")
r.font.name = 'Times New Roman'; r.font.size = Pt(18)

p2 = doc.add_paragraph()
set_para_format(p2, align='center', space_before=6, space_after=0)
r2 = p2.add_run("DIPLOMA KEMAHIRAN MALAYSIA")
r2.font.name = 'Times New Roman'; r2.font.size = Pt(18)

add_blank(3)

p3 = doc.add_paragraph()
set_para_format(p3, align='center', space_before=0, space_after=0)
r3 = p3.add_run("I DEFENDER")
r3.font.name = 'Times New Roman'; r3.font.size = Pt(18); r3.font.bold = True

add_blank(4)

p4 = doc.add_paragraph()
set_para_format(p4, align='center', space_before=0, space_after=0)
r4 = p4.add_run("[NAMA PELATIH 1]\n[NAMA PELATIH 2]\n[NAMA PELATIH 3]")
r4.font.name = 'Times New Roman'; r4.font.size = Pt(18)

add_blank(5)

p5 = doc.add_paragraph()
set_para_format(p5, align='center', space_before=0, space_after=0)
r5 = p5.add_run("[NAMA PROGRAM]\n[KOD PROGRAM]")
r5.font.name = 'Times New Roman'; r5.font.size = Pt(18)

doc.add_page_break()

# ════════════════════════════════════════════════
#  ABSTRAK
# ════════════════════════════════════════════════
p_abs_title = doc.add_paragraph()
set_para_format(p_abs_title, align='center', space_before=0, space_after=18)
r = p_abs_title.add_run("ABSTRAK")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

add_body(
    "Projek ini bertajuk I Defender merupakan sebuah sistem pengesanan pencerobohan fizikal secara masa nyata (real-time physical intrusion detection system) yang dibangunkan menggunakan Raspberry Pi 4 Model B 8GB RAM, Raspberry Pi Camera Module 3 (Camera Module V3), perpustakaan pemprosesan imej OpenCV dan model pengesanan objek YOLO (You Only Look Once). Sistem ini direka untuk memantau sesebuah kawasan secara automatik, mengesan kehadiran manusia yang tidak dibenarkan dan mengeluarkan amaran secara serta-merta tanpa memerlukan pemantauan manual oleh manusia.",
    tab=True
)
add_blank()
add_body(
    "Metodologi pembangunan projek ini menggunakan Model Air Terjun (Waterfall Model) yang merangkumi fasa perancangan, analisis, rekabentuk, implementasi dan pengujian. Perkakasan utama yang digunakan ialah Raspberry Pi 4 Model B 8GB, Camera Module V3, kad MicroSD 64GB dan bekalan kuasa USB-C 5V/3A. Perisian yang digunakan pula ialah Raspberry Pi OS 64-bit, Python 3, OpenCV 4, model YOLOv8 Nano serta perpustakaan Ultralytics sebagai rangka kerja pengesanan objek.",
    tab=True
)
add_blank()
add_body(
    "Hasil projek menunjukkan bahawa sistem I Defender berjaya mengesan kehadiran manusia dalam kawasan yang dipantau dengan kadar ketepatan yang memuaskan dalam persekitaran makmal yang terkawal. Sistem ini berupaya memproses aliran video secara langsung (live stream), melukis kotak sempadan (bounding box) di sekeliling objek yang dikesan dan mencetuskan amaran apabila pencerobohan dikesan. Kesimpulannya, I Defender merupakan penyelesaian pemantauan keselamatan yang praktikal dan kos efektif, sesuai untuk kegunaan institusi kecil, pejabat dan premis yang memerlukan pengawasan automatik.",
    tab=True
)

doc.add_page_break()

# ════════════════════════════════════════════════
#  BAB I
# ════════════════════════════════════════════════
add_heading_bab("BAB I", "PENGENALAN")

add_section_heading("1.1", "PENGENALAN PROJEK")
add_body(
    "Projek I Defender merupakan sebuah sistem kamera keselamatan pintar yang mampu mengesan kehadiran manusia secara automatik menggunakan teknologi penglihatan komputer (computer vision). Sistem ini dibangunkan di atas platform Raspberry Pi 4 Model B 8GB RAM yang dilengkapi dengan Raspberry Pi Camera Module V3 sebagai peranti penangkapan imej. Pemprosesan video dijalankan menggunakan perpustakaan OpenCV manakala pengesanan objek khususnya pengesanan manusia dilakukan menggunakan model YOLOv8 Nano yang dioptimumkan untuk berjalan pada perkakasan berdaya rendah.",
    tab=True
)
add_blank()
add_body(
    "Sistem perakam keselamatan (CCTV) konvensional yang sedia ada hanya berfungsi sebagai peranti rakaman pasif yang memerlukan pemantauan manusia secara berterusan untuk mengesan sebarang kejadian yang mencurigakan. Pendekatan ini tidak cekap dan tidak praktikal terutamanya bagi premis yang tidak mempunyai kakitangan keselamatan yang mencukupi. I Defender hadir sebagai penyelesaian kepada masalah ini dengan menggabungkan keupayaan rakaman video bersama pengesanan objek secara automatik dalam satu sistem yang padat, berkos rendah dan mudah dipasang. Jangkaan hasil projek ini ialah sebuah sistem yang dapat beroperasi secara berterusan selama 24 jam, mengesan kehadiran manusia dalam masa kurang dari satu saat dan menghantar amaran kepada pentadbir apabila pencerobohan dikesan.",
    tab=True
)

add_blank()
add_section_heading("1.2", "LATARBELAKANG MASALAH")
add_body(
    "Keselamatan premis merupakan satu keperluan asas bagi setiap organisasi, institusi pendidikan dan perniagaan. Sistem kamera keselamatan (CCTV) konvensional telah lama digunakan sebagai kaedah pengawasan, namun sistem ini mempunyai beberapa kelemahan yang ketara. Kajian dan pemerhatian yang dijalankan mendapati bahawa kebanyakan sistem CCTV yang dipasang hanya berfungsi sebagai alat rakaman pasif dan tidak mampu memberikan tindak balas atau amaran secara automatik apabila berlaku pencerobohan.",
    tab=True
)
add_blank()
add_body(
    "Masalah pertama yang dikenal pasti ialah ketidakupayaan sistem CCTV sedia ada untuk mengesan pencerobohan secara automatik. Pemantauan rakaman CCTV memerlukan kehadiran manusia secara berterusan, iaitu satu keperluan yang mahal dan tidak praktikal. Masalah kedua ialah sistem penggera keselamatan konvensional yang berasaskan penderia gerakan (PIR sensor) tidak mampu membezakan antara manusia, haiwan atau objek bergerak lain, menyebabkan berlakunya amaran palsu (false alarm) yang kerap. Masalah ketiga ialah kos sistem kamera keselamatan pintar komersial yang dilengkapi keupayaan pengesanan automatik adalah sangat tinggi, menjadikannya di luar kemampuan institusi kecil dan perniagaan berskala kecil.",
    tab=True
)
add_blank()
add_body(
    "Berdasarkan masalah-masalah yang dikenal pasti ini, satu keperluan jelas wujud untuk membangunkan sistem pemantauan keselamatan yang berupaya mengesan kehadiran manusia secara automatik, membezakan manusia daripada objek lain, memberikan amaran serta-merta dan boleh dibangunkan dengan kos yang rendah. Projek I Defender dibangunkan untuk memenuhi keperluan ini dengan memanfaatkan perkakasan Raspberry Pi 4 8GB dan model pengesanan objek YOLO yang berjalan di atas perpustakaan OpenCV.",
    tab=True
)

add_blank()
add_section_heading("1.3", "OBJEKTIF PROJEK")
add_body("Objektif kajian adalah seperti berikut:", tab=True)
add_blank()
add_list_item("i)", "Membangunkan sistem pengesanan kehadiran manusia secara masa nyata menggunakan Raspberry Pi 4 Model B 8GB, Raspberry Pi Camera Module V3, perpustakaan OpenCV 4 dan model pengesanan objek YOLOv8 Nano yang mampu memproses aliran video langsung dengan kadar bingkai (frame rate) yang mencukupi untuk tujuan pemantauan keselamatan.")
add_blank()
add_list_item("ii)", "Mengkonfigurasi model YOLOv8 agar mampu mengesan kehadiran manusia dengan tepat dalam pelbagai keadaan pencahayaan dan sudut pandangan dalam persekitaran makmal, serta memaparkan kotak sempadan (bounding box) dan paras keyakinan (confidence score) pada setiap objek yang dikesan.")
add_blank()
add_list_item("iii)", "Membangunkan modul amaran automatik yang akan mencetuskan pemberitahuan kepada pentadbir sistem apabila kehadiran manusia dikesan di dalam kawasan yang dipantau, termasuk merekodkan tangkapan gambar (snapshot) berserta cap masa (timestamp) sebagai bukti kejadian.")

add_blank()
add_section_heading("1.4", "SKOP PROJEK")
add_body("Bagi mencapai matlamat dan objektif yang digariskan, beberapa skop projek telah dikenalpasti.", tab=True)
add_blank()
add_body("Antaranya ialah:", tab=True)
add_blank()
add_list_item("i.", "Projek ini meliputi pembangunan, konfigurasi dan pengujian sistem I Defender untuk memantau satu kawasan dalaman (indoor) sahaja, iaitu kawasan makmal komputer di E-Access International College, dengan satu unit kamera Camera Module V3 yang dipasang pada satu sudut pandangan tetap.")
add_blank()
add_list_item("ii.", "Batasan dari segi pengesanan objek adalah tertumpu kepada pengesanan kelas manusia (person class) sahaja berdasarkan set data model YOLOv8, dan tidak merangkumi pengesanan objek lain seperti kenderaan, haiwan atau barang.")
add_blank()
add_list_item("iii.", "Batasan dari segi tempoh adalah selama empat (4) bulan, bermula dari bulan Januari hingga April, yang merangkumi keseluruhan fasa perancangan, pembangunan, pengujian dan dokumentasi projek.")
add_blank()
add_list_item("iv.", "Sistem I Defender dibangunkan menggunakan bahasa pengaturcaraan Python 3 dan tidak merangkumi pembangunan aplikasi mudah alih (mobile application) atau antaramuka web yang kompleks.")
add_blank()
add_list_item("v.", "Pengujian sistem dijalankan dalam persekitaran makmal yang terkawal dengan keadaan pencahayaan biasa (normal indoor lighting) dan tidak merangkumi pengujian dalam keadaan cahaya rendah (low light) atau persekitaran luar (outdoor).")

doc.add_page_break()

# ════════════════════════════════════════════════
#  BAB II
# ════════════════════════════════════════════════
add_heading_bab("BAB II", "METODOLOGI")

add_section_heading("2.1", "PENGENALAN")
add_body(
    "Bab ini akan menerangkan secara terperinci kaedah-kaedah yang digunakan dan aliran keseluruhan projek I Defender dari peringkat perancangan sehinggalah pengujian. Penerangan merangkumi susun atur (set-up) perkakasan iaitu Raspberry Pi 4 Model B 8GB dan Camera Module V3, konfigurasi persekitaran perisian Python, OpenCV dan YOLOv8, serta aliran pemprosesan video dan pengesanan objek yang menjadi teras kepada fungsi sistem ini. Langkah-langkah pengubahsuaian (troubleshooting) yang diambil semasa fasa implementasi turut diterangkan bagi memberikan gambaran yang lengkap tentang proses pembangunan sistem.",
    tab=True
)

add_blank()
add_section_heading("2.2", "SUSUN ATUR PROJEK")
add_body(
    "Projek I Defender menggunakan Model Air Terjun (Waterfall Model) sebagai rangka kerja pembangunan. Pendekatan ini dipilih kerana keperluan sistem telah dikenal pasti dengan jelas sejak awal, dan setiap fasa pembangunan mempunyai hasil (deliverable) yang boleh disahkan sebelum fasa seterusnya dimulakan. Model ini juga memudahkan pengurusan masa dan sumber dalam tempoh empat bulan yang ditetapkan.",
    tab=True
)

add_blank()
add_subheading("Fasa 1: Perancangan (Planning)")
add_body(
    "Pada fasa ini, skop projek ditetapkan, keperluan perkakasan dan perisian dikenal pasti, dan jadual pelaksanaan dalam bentuk Carta Gantt disediakan. Sesi perbincangan bersama penyelia dijalankan untuk memastikan hala tuju projek adalah tepat. Kajian awal terhadap keupayaan Raspberry Pi 4 dalam memproses video masa nyata dan menjalankan model YOLO turut dilaksanakan pada fasa ini bagi memastikan perkakasan yang dipilih mampu memenuhi keperluan prestasi sistem.",
    tab=True
)

add_blank()
add_subheading("Fasa 2: Analisis Sistem (System Analysis)")
add_body(
    "Fasa analisis melibatkan kajian mendalam terhadap teknologi pengesanan objek yang sesuai untuk digunakan pada platform Raspberry Pi 4. Kajian perbandingan antara pelbagai versi model YOLO iaitu YOLOv5, YOLOv7 dan YOLOv8 dijalankan dari segi ketepatan pengesanan (detection accuracy), kelajuan pemprosesan (inference speed) dan penggunaan memori pada seni bina ARM Cortex-A72 milik Raspberry Pi 4. Hasil analisis mendapati YOLOv8 Nano (YOLOv8n) adalah pilihan terbaik kerana menawarkan keseimbangan yang baik antara kelajuan dan ketepatan pengesanan pada perkakasan berdaya rendah. Camera Module V3 pula dipilih sebagai kamera kerana menyokong resolusi sehingga 12 megapiksel dan mempunyai antara muka CSI (Camera Serial Interface) yang terus disambungkan ke Raspberry Pi 4 tanpa memerlukan pemacu tambahan.",
    tab=True
)

add_blank()
add_subheading("Fasa 3: Rekabentuk Sistem (System Design)")
add_body(
    "Rekabentuk sistem I Defender merangkumi tiga komponen utama iaitu komponen penangkapan video, komponen pemprosesan dan pengesanan, serta komponen amaran. Aliran kerja sistem adalah seperti berikut: Camera Module V3 menangkap aliran video secara berterusan, setiap bingkai video (frame) diterima oleh OpenCV untuk pra-pemprosesan (preprocessing) termasuk pengubahsuaian saiz (resizing) dan penukaran ruang warna (color space conversion), bingkai yang telah diproses dihantar kepada model YOLOv8n untuk pengesanan objek, hasil pengesanan berupa koordinat kotak sempadan (bounding box coordinates) dan nilai keyakinan (confidence value) dikembalikan, dan akhirnya apabila manusia dikesan dengan nilai keyakinan melebihi ambang yang ditetapkan (confidence threshold), sistem akan mencetuskan amaran dan merekodkan tangkapan gambar berserta cap masa.",
    tab=True
)

add_blank()
add_subheading("Fasa 4: Implementasi (Implementation)")
add_body(
    "Pelaksanaan implementasi dimulakan dengan pemasangan sistem operasi Raspberry Pi OS 64-bit (Bookworm) pada kad MicroSD 64GB menggunakan perisian Raspberry Pi Imager. Setelah sistem operasi berjalan, persekitaran Python 3.11 dikonfigurasikan dan semua pakej yang diperlukan dipasang menggunakan pengurus pakej pip, termasuk OpenCV-Python, perpustakaan Ultralytics untuk YOLOv8 dan pakej picamera2 untuk mengawal Camera Module V3. Model YOLOv8n dimuat turun dan diuji pada peringkat awal menggunakan beberapa imej ujian bagi memastikan pengesanan berfungsi dengan betul sebelum integrasi penuh dengan aliran video langsung dijalankan. Kod pengaturcaraan utama sistem ditulis dalam bahasa Python menggunakan pendekatan pengaturcaraan berorientasikan objek (OOP) bagi memudahkan penyelenggaraan dan pengubahsuaian pada masa hadapan.",
    tab=True
)

add_blank()
add_subheading("Fasa 5: Pengujian (Testing)")
add_body(
    "Pengujian sistem dijalankan melalui beberapa senario yang berbeza. Ujian pertama ialah ujian pengesanan statik (static detection test) di mana beberapa orang individu diminta berdiri di hadapan kamera pada jarak yang berbeza-beza untuk mengesahkan keupayaan sistem mengesan manusia dari pelbagai jarak. Ujian kedua ialah ujian pengesanan dinamik (dynamic detection test) di mana individu diminta bergerak melintas di hadapan kamera untuk menguji keupayaan sistem mengesan gerakan. Ketepatan pengesanan, kadar amaran palsu (false positive rate) dan kadar bingkai pemprosesan (frames per second / FPS) direkodkan dan dianalisis bagi setiap senario ujian.",
    tab=True
)

add_blank()

# Hardware table
p_hw = doc.add_paragraph()
set_para_format(p_hw, align='left', space_before=6, space_after=4)
r = p_hw.add_run("Keperluan Perkakasan:")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

add_table(
    ["Bil", "Perkakasan", "Spesifikasi", "Kuantiti"],
    [
        ["1", "Raspberry Pi 4 Model B", "8GB LPDDR4X RAM, ARM Cortex-A72 @ 1.8GHz", "1 unit"],
        ["2", "Raspberry Pi Camera Module V3", "12MP Sony IMX708, Sudut Pandangan 66°", "1 unit"],
        ["3", "Kad MicroSD", "64GB, Kelas 10 / A2, Kelajuan Baca ≥ 90MB/s", "1 keping"],
        ["4", "Bekalan Kuasa USB-C", "5V / 3A (15W)", "1 unit"],
        ["5", "Kabel CSI", "15cm, sambungan Camera Module ke Pi", "1 utas"],
        ["6", "Casing Raspberry Pi", "Dengan soket kamera dan penyejuk udara", "1 unit"],
        ["7", "Monitor / Skrin", "Untuk paparan semasa pengujian", "1 unit"],
        ["8", "Komputer Riba", "Untuk pembangunan kod dan pemindahan fail", "1 unit"],
    ]
)

p_sw = doc.add_paragraph()
set_para_format(p_sw, align='left', space_before=6, space_after=4)
r = p_sw.add_run("Keperluan Perisian:")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

add_table(
    ["Bil", "Perisian / Pustaka", "Versi", "Fungsi"],
    [
        ["1", "Raspberry Pi OS 64-bit", "Bookworm (Debian 12)", "Sistem Operasi Utama"],
        ["2", "Python", "3.11", "Bahasa Pengaturcaraan Utama"],
        ["3", "OpenCV-Python", "4.8.x", "Pemprosesan dan Analisis Video"],
        ["4", "Ultralytics (YOLOv8)", "8.x", "Rangka Kerja Pengesanan Objek"],
        ["5", "YOLOv8n Model", "yolov8n.pt", "Model Pengesanan Objek (Nano)"],
        ["6", "picamera2", "0.3.x", "Kawalan Raspberry Pi Camera Module V3"],
        ["7", "NumPy", "1.24.x", "Operasi Tatasusunan dan Matriks"],
        ["8", "Thonny IDE", "4.x", "Persekitaran Pembangunan Kod"],
    ]
)

add_blank()
add_section_heading("2.3", "KESIMPULAN")
add_body(
    "Secara keseluruhannya, Bab II ini telah menerangkan dengan terperinci metodologi dan susun atur yang digunakan dalam pembangunan projek I Defender. Pemilihan Raspberry Pi 4 Model B 8GB sebagai platform pemprosesan utama terbukti mampu menampung keperluan pengesanan objek masa nyata menggunakan YOLOv8 Nano dengan prestasi yang memuaskan. Kamera Camera Module V3 pula menyediakan kualiti imej yang tinggi dan integrasi yang lancar dengan platform Raspberry Pi, manakala gabungan OpenCV dan YOLOv8 membentuk saluran pemprosesan (processing pipeline) yang cekap untuk tujuan pengesanan kehadiran manusia.",
    tab=True
)
add_blank()
add_body(
    "Pemilihan Model Air Terjun (Waterfall Model) sebagai rangka kerja pembangunan memastikan setiap fasa dilaksanakan secara sistematik dan teratur dalam tempoh empat bulan yang ditetapkan. Susun atur perkakasan yang padat, penggunaan perisian open-source yang terbukti berkesan, serta prosedur pengujian yang komprehensif merupakan faktor-faktor utama yang menyumbang kepada kelancaran pembangunan sistem I Defender. Hasil implementasi dan analisis keputusan pengujian akan dibincangkan dengan lebih lanjut dalam bab seterusnya.",
    tab=True
)

# ── Save ──
out_path = "/home/user/dataai-Landing/Laporan_I_Defender_DKM.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
